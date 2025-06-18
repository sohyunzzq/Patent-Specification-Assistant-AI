from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt,RGBColor,Parented
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.oxml import parse_xml
import pandas as pd
import re

import torch
import nltk
from transformers import AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig, AutoModelForSeq2SeqLM
from peft import PeftModel, PeftConfig
from peft import prepare_model_for_kbit_training, LoraConfig, get_peft_model

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.linear_model import SGDClassifier
from okt import build_bag_of_words
import numpy as np
import json

import requests
import xmltodict

from transformers import BertTokenizer, BertPreTrainedModel, AdamW, BertConfig, BertModel
import torch.nn as nn
import pickle

#//////////////////////////////////////////////////////
_abc = 'C'
text = input("\n\n발명의 명칭 : ")
text1 = input("\n\n발명의 구성 : ")
text_effect = input("\n\n발명의 효과 : ")
document=Document()
style=document.styles['Normal']
font=style.font
font.name='나눔고딕'
style.font.size=Pt(10)
document_name="NCCOSS_KOSPAI.docx"
#//////////////////////////////////////////////////////
paragraph= document.add_paragraph()
temp = []

def gen(x):
    q = f"### 질문: {x}\n\n### 답변:"
    gened = model.generate(
        **tokenizer(
            q, 
            return_tensors='pt', 
            return_token_type_ids=False
        ).to('cuda'), 
        max_new_tokens=256,
        early_stopping=True,
        do_sample=True,
        eos_token_id=2,
    )
    generated_text = tokenizer.decode(gened[0], skip_special_tokens=True)
    answer_start_index = generated_text.find("### 답변:") + len("### 답변:")
    generated_answer = generated_text[answer_start_index:].strip()
    d= paragraph.add_run(generated_answer)
    paragraph.add_run('\n')
    feild = paragraph.add_run('_'*79)
    feild.bold=True
    return generated_answer


def gen_solution(x):
    q = f"### 질문: {x}\n\n### 답변:"
    gened = model.generate(
        **tokenizer(
            q, 
            return_tensors='pt', 
            return_token_type_ids=False
        ).to('cuda'), 
        max_new_tokens=512,
        early_stopping=True,
        do_sample=True,
        eos_token_id=2,
    )
    generated_text = tokenizer.decode(gened[0], skip_special_tokens=True)
    answer_start_index = generated_text.find("### 답변:") + len("### 답변:")
    generated_answer = generated_text[answer_start_index:].strip()
    d= paragraph.add_run(generated_answer)
    paragraph.add_run('\n')
    feild = paragraph.add_run('_'*79)
    feild.bold=True
    return generated_answer

#/////////////////////////////////////////////////////

paragraph.add_run(f"발명의 명칭 ")
text_name = re.sub(r"[^가-힣\s]", "", text)
prior_name=paragraph.add_run(text_name)
prior_name.bold=True
paragraph.add_run('\n')
feild = paragraph.add_run('-'*119)
paragraph.add_run('\n')
feild.bold=True
feild.underline=True
feild.italic=True

#//////////////////////////////////////////////////////

prior_art=paragraph.add_run("[선행기술문헌]\n")
prior_art.font.color.rgb=RGBColor(255,0,127)
prior_art.bold=True
prior_art.italic=True
input1 = text

df = pd.read_csv('/home/in2wise/economy/keyword.csv', encoding='UTF8')
consider = ['발명의 명칭','keyword']

df = df[consider]
df = df.astype('str')
input_df = df.fillna('')
inputtt = input_df[['keyword']]
docs = inputtt['keyword'].values.tolist()
len_docs = len(docs)

with open('/home/in2wise/NCCOSS/stopwords.json', 'r') as f:
    stop_list = json.load(f)
f.close()

test2 = []
for words in input1.split():
    vocab, bow = build_bag_of_words(words)
    test = []
    for word in vocab:
        if word in stop_list:
            continue
        test.append(word)
    test2.append(''.join(test))
docs.append(' '.join(test2))

tfidfv = TfidfVectorizer().fit(docs)
X_tfidf = tfidfv.fit_transform(docs)

feature_names = tfidfv.get_feature_names_out()
tfidf_df = pd.DataFrame(X_tfidf.toarray(), columns=feature_names)
np.set_printoptions()

key = []
for doc_index, doc_row in enumerate(X_tfidf):
    non_zero_elements = doc_row.nonzero()[1]
    if non_zero_elements.any():
        words = {}
        for word_index in non_zero_elements:
            word = list(tfidfv.vocabulary_.keys())[list(tfidfv.vocabulary_.values()).index(word_index)]
            words[word] = doc_row[0, word_index]
        words = sorted(words.items(), key=lambda x: x[1], reverse=True)
        tmp = []
        for k in range(len(words) if len(words) < 2 else 2):
            tmp.append(words[k][0])
    key.append(' '.join(tmp))

###### TF-IDF로 키워드 추출 완료

url1 = "http://plus.kipris.or.kr/kipo-api/kipi/patUtiModInfoSearchSevice/getAdvancedSearch?inventionTitle="
url2 = "&ServiceKey="
RESTkey= "VuIvh4dUvFerjHsiaEudNFt1AKBpEiBm=bVJNBYbnI8=" #키 입력

key = key[len_docs].replace(' ', '*')

reponse = requests.get(url1+key+url2+RESTkey)
content = reponse.content
dict_type = xmltodict.parse(content)
items = dict_type['response']['body']['items']

i = 1    
if items == None:
    final = "\n"
elif isinstance(items['item'], list):
    for name in items['item']:
        if i == 1:            
            final = "(특허문헌 000" + str(i) + ") 한국등록특허 제10-" + str(name['applicationNumber'])[2:6] + '-' + str(name['applicationNumber'])[6:] + "호 <" + name['inventionTitle'] + ">"
        else:
            final = final + "\n(특허문헌 000" + str(i) + ") 한국등록특허 제10-" + str(name['applicationNumber'])[2:6] + '-' + str(name['applicationNumber'])[6:] + "호 <" + name['inventionTitle'] + ">"
        i += 1
        if i > 3:
            break
elif isinstance(items['item'], dict):  # Only one item
    final = "(특허문헌 0001) 한국등록특허 제10-" + str(items['item']['applicationNumber'])[2:6] + '-' + str(items['item']['applicationNumber'])[6:] + "호 <" + items['item']['inventionTitle'] + ">"
else:
    for name in items['item']:
        if i == 1:            
            final = "(특허문헌 000" + str(i) + ") 한국등록특허 제10-" + str(name['applicationNumber'])[2:6] + '-' + str(name['applicationNumber'])[6:] + "호 <" + name['inventionTitle'] + ">"
paragraph.add_run(final)
paragraph.add_run('\n')
feild = paragraph.add_run('_'*79)
print('\n\n\n 선행기술문헌 작성 완료\n\n\n')

#//////////////////////////////////////////////////////

field=paragraph.add_run("\n[기 술 분 야]\n\n")
field.font.color.rgb=RGBColor(255,0,127)
field.bold=True
field.italic=True
peft_model_id = "/home/in2wise/mis_기술분야_epochs15"
config = PeftConfig.from_pretrained(peft_model_id)

bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_use_double_quant=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.bfloat16
)

model = AutoModelForCausalLM.from_pretrained(
    config.base_model_name_or_path,
    quantization_config=bnb_config,
    device_map={"": 0}
)

model = PeftModel.from_pretrained(model, peft_model_id)

tokenizer = AutoTokenizer.from_pretrained(config.base_model_name_or_path)
model.eval()


temp.append(gen(text))
paragraph.add_run("\n")
print('\n\n\n 기술분야 작성 완료\n\n\n')

#//////////////////////////////////////////////////////

background=paragraph.add_run("\n[배 경 기 술]\n\n")
background.font.color.rgb=RGBColor(255,0,127)
background.bold=True
background.italic=True
peft_model_id = "/home/in2wise/배경기술_fin"
config = PeftConfig.from_pretrained(peft_model_id)

bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_use_double_quant=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.bfloat16
)

model = AutoModelForCausalLM.from_pretrained(
    config.base_model_name_or_path,
    quantization_config=bnb_config,
    device_map={"": 0}
)

model = PeftModel.from_pretrained(model, peft_model_id)

tokenizer = AutoTokenizer.from_pretrained(config.base_model_name_or_path)
model.eval()


text_tec = gen(text)
print('\n\n\n 배경기술 작성 완료\n\n\n')

#//////////////////////////////////////////////////////

boundary=paragraph.add_run("\n[청 구 범 위]\n\n")
boundary.font.color.rgb=RGBColor(255,0,127)
boundary.bold=True
boundary.italic=True
paragraph.add_run(text1)
paragraph.add_run('\n')
feild = paragraph.add_run('_'*79)
feild.bold=True

#//////////////////////////////////////////////////////

solve=paragraph.add_run("\n[해결하려는 과제]\n\n")
solve.font.color.rgb=RGBColor(255,0,127)
solve.bold=True
solve.italic=True

peft_model_id = "/home/in2wise/economy/KoAlpaca_vol13_save"
config = PeftConfig.from_pretrained(peft_model_id)

bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_use_double_quant=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.bfloat16
)

model = AutoModelForCausalLM.from_pretrained(
    config.base_model_name_or_path,
    quantization_config=bnb_config,
    device_map={"": 0}
)

model = PeftModel.from_pretrained(model, peft_model_id)

tokenizer = AutoTokenizer.from_pretrained(config.base_model_name_or_path)
model.eval()
gen(text)

print('\n\n\n 해결하려는 과제 작성 완료\n\n\n')

#//////////////////////////////////////////////////////

solve_tools=paragraph.add_run("\n[과제 해결 수단]\n\n")
solve_tools.font.color.rgb=RGBColor(255,0,127)
solve_tools.bold=True
solve_tools.italic=True

peft_model_id = "/home/in2wise/solution_epochs20"
config = PeftConfig.from_pretrained(peft_model_id)

bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_use_double_quant=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.bfloat16
)

model = AutoModelForCausalLM.from_pretrained(
    config.base_model_name_or_path,
    quantization_config=bnb_config,
    device_map={"": 0}
)

model = PeftModel.from_pretrained(model, peft_model_id)

tokenizer = AutoTokenizer.from_pretrained(config.base_model_name_or_path)
model.eval()



text_solution = gen_solution(text1)


print('\n\n\n 과제의 해결 수단 작성 완료\n\n\n')

#//////////////////////////////////////////////////////

effect=paragraph.add_run("\n[발명의 효과]\n\n")
effect.font.color.rgb=RGBColor(255,0,127)
effect.bold=True
effect.italic=True
paragraph.add_run(text_effect)
paragraph.add_run('\n')
feild = paragraph.add_run('_'*79)
feild.bold=True
temp.append(text_effect)

#//////////////////////////////////////////////////////

str1 = '\n'.join(temp)
summarize=paragraph.add_run("\n[요약]\n\n")
summarize.font.color.rgb=RGBColor(255,0,127)
summarize.font.size=Pt(10)
summarize.bold=True
summarize.italic=True

nltk.download('punkt')

model_dir = "lcw99/t5-large-korean-text-summary"
tokenizer = AutoTokenizer.from_pretrained(model_dir)
model = AutoModelForSeq2SeqLM.from_pretrained(model_dir)

max_input_length = 512 + 256

# text = str1
inputs = [text_tec]

inputs = tokenizer(inputs, max_length=max_input_length, truncation=True, return_tensors="pt")
output = model.generate(**inputs, num_beams=8, do_sample=True, min_length=10, max_length=100)
decoded_output = tokenizer.batch_decode(output, skip_special_tokens=True)[0]
predicted_title = nltk.sent_tokenize(decoded_output.strip())[0]
paragraph.add_run(predicted_title)
#print(predicted_title)


# text2 = input.at[index,'과제의 해결수단']
# inputs = [text2]
inputs = [text_effect]
inputs = tokenizer(inputs, max_length=max_input_length, truncation=True, return_tensors="pt")
output = model.generate(**inputs, num_beams=8, do_sample=True, min_length=10, max_length=100)
decoded_output = tokenizer.batch_decode(output, skip_special_tokens=True)[0]
predicted_title = nltk.sent_tokenize(decoded_output.strip())[0]
paragraph.add_run(predicted_title)
#print(predicted_title)


# text3 = input.at[index,'발명의 효과']
# inputs = [text3]
inputs = [text_solution]
inputs = tokenizer(inputs, max_length=max_input_length, truncation=True, return_tensors="pt")
output = model.generate(**inputs, num_beams=8, do_sample=True, min_length=10, max_length=100)
decoded_output = tokenizer.batch_decode(output, skip_special_tokens=True)[0]
predicted_title = nltk.sent_tokenize(decoded_output.strip())[0]
paragraph.add_run(predicted_title)
#print(predicted_title)
feild = paragraph.add_run('_'*79)
feild.bold=True


print('\n\n\n 요약 작성 완료')

#////////////////////////////////////////////////////////////////

class BertForMultiLabelSequenceClassification(BertPreTrainedModel):
    def __init__(self, config):
        super(BertForMultiLabelSequenceClassification, self).__init__(config)
        self.num_labels = config.num_labels

        self.bert = BertModel(config)
        self.dropout = nn.Dropout(0.5)
        self.classifier = nn.Linear(config.hidden_size, config.num_labels)

        self.init_weights()

    def forward(
        self,
        input_ids=None,
        attention_mask=None,
        token_type_ids=None,
        position_ids=None,
        head_mask=None,
        inputs_embeds=None,
        labels=None,
    ):
        _, pooled_output = self.bert(input_ids, attention_mask)

        pooled_output = self.dropout(pooled_output)
        logits = self.classifier(pooled_output)

        return logits

    def freeze_bert_encoder(self):
        for param in self.bert.parameters():
            param.requires_grad = False
    
    def unfreeze_bert_encoder(self):
        for param in self.bert.parameters():
            param.requires_grad = True

from keras.preprocessing.sequence import pad_sequences

def text_to_loader_bert(sentences, max_len):
    #tokenizer = AutoTokenizer.from_pretrained("skt/kobert-base-v1")
    tokenizer = AutoTokenizer.from_pretrained("monologg/koelectra-base-v3-discriminator")
    
    sentences = [tokenizer.tokenize(sentence) for sentence in sentences]
    [sentence.insert(0,'[CLS]') for sentence in sentences]
    [sentence.append('[SEP]') if len(sentence)<max_len-1 else sentence.insert(max_len-1, '[SEP]') for sentence in sentences]
    
    input_ids = [tokenizer.convert_tokens_to_ids(sentence) for sentence in sentences]
    input_ids = pad_sequences(input_ids, maxlen=max_len, dtype='long', padding='post', truncating='post')
    
    attention_masks = []
    
    for input_id in input_ids:
        attention_mask = [1 if (i>0) else 0 for i in input_id]
        attention_masks.append(attention_mask)
    
    #mlb = MultiLabelBinarizer()
    #mlb.fit(labels)
    #labels = mlb.transform(labels)
    
    train_inputs = torch.tensor(input_ids)
    train_masks = torch.tensor(attention_masks)
    # train_labels = torch.tensor(labels)
    
    return train_inputs, train_masks


test_input, test_mask = text_to_loader_bert(text1, 512)
test_input.to("cuda")
test_mask.to("cuda")

model = BertForMultiLabelSequenceClassification.from_pretrained('/home/in2wise/classification/MultiLabel/model/bert_D', cache_dir=None, num_labels=8)
model = model.cuda()

with open('/home/in2wise/classification/MultiLabel/tensor_D/labels_le.pkl', 'rb') as f:
        le = pickle.load(f)
with open('/home/in2wise/classification/MultiLabel/tensor_D/labels_lb.pkl', 'rb') as f:
        lb = pickle.load(f)

arr1 = []
arr2 = []

test = model(test_input[0:1].to("cuda"), token_type_ids=None, attention_mask=test_mask[0:1].to("cuda"))
values, indexs = torch.topk(test,2)

for index in indexs:
    arr1.append(index[0].item())
    arr2.append(index[1].item())

result1 = le.inverse_transform(arr1)
result2 = le.inverse_transform(arr2)

result1.tolist()
result2.tolist()
result1 = list(result1[0])
result2 = list(result2[0])
result1[0] = _abc
result2[0] = _abc
result1 = ''.join(result1)
result2 = ''.join(result2)

summarize=paragraph.add_run("\n[분류]\n\n")
summarize.font.color.rgb=RGBColor(255,0,127)
summarize.font.size=Pt(10)
summarize.bold=True
summarize.italic=True

paragraph.add_run(result1)
paragraph.add_run('\n')
paragraph.add_run(result2)

print('\n\n\n\n 특허명세서 작성 완료')

document.save(document_name)